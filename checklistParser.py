from docx import Document
from xmlObject import XMLObject
from resources import Resources
import xml.etree.ElementTree as ET
import const
import logging
import wx

def isWordTypeTable(child):
    if child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_TABLE:
        return True
    else:
        return False

def isWordTypeCheckboxlist(child):
    if child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_CHECKBOXLIST:
        return True
    else:
        return False

def isWordTypeYesNolist(child):
    if child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_YESNOLIST:
        return True
    else:
        return False

def readVersion(checklistDocument):
    log = logging.getLogger("DSEGenerator.readVersion")
    try:
        version = checklistDocument.tables[0].table.cell(1, 0).text
        print("version: " + version)
        if version is None:
            version = "1.0"
        return version
    except (IndexError):
        log.error("list index out of range!")
        wx.MessageBox("Can't determine version of selected Document! Processing aborted. Please select another one!", caption="Error occured!")

def parseYesNoList(xmlElement, checklistDocument):
    log = logging.getLogger("DSEGenerator.parseYesNoList")
    checkboxStatelist = []
    element = checklistDocument._element.xml
    elementliste = element.split("<")
    log.debug("doc: " + str(checklistDocument.text))
    for e in elementliste:
        log.debug("element: " + str(e))
        if e.startswith("w14:checked "):
            e = e.split('"')[1]
            if e == "1" : 
                checkboxStatelist.append(True)
            else:
                checkboxStatelist.append(False)
    if checkboxStatelist is not None and checkboxStatelist[0] is True:
        return "True"
    else:
        return "False"

def parseCheckboxlist(xmlElement, checklistDocument):
    log = logging.getLogger("DSEGenerator.parseCheckboxlist")
    checkboxStatelist = []
    values = {}
    i = 0    
    element = checklistDocument._element.xml
    elementliste = element.split("<")
    for e in elementliste:
        if e.startswith("w14:checked "):
            e = e.split('"')[1]
            if e == "1" : 
                checkboxStatelist.append(True)
            else:
                checkboxStatelist.append(False)
    log.debug("list: " + str(checkboxStatelist))
    log.debug("doc: " + str(checklistDocument.text))
    for e in checklistDocument.text.split("\n"):
        e=e.strip()
        if i==0:
            i=i+1
            continue
        else:
            try:
                if len(e) > 0:
                    log.debug("append key: " + e + "value: " + str(checkboxStatelist[i-1]) )
                    values[e] = checkboxStatelist[i-1]
            except (IndexError):
                log.error("list index out of range!")
        i=i+1
    log.debug("returning values: " + str(values))
    return values

def parseTable(xmlElement, checklistDocument):
    pos = int(xmlElement.attrib.get(const.CHECKLIST_ATTRIB_TAB))
    table = checklistDocument.tables[pos-1]
    tableObject = {}    
    for child in xmlElement:
        zeile = int(child.attrib.get(const.CHECKLIST_ATTRIB_ROW))-1
        spalte = int(child.attrib.get(const.CHECKLIST_ATTRIB_COL))-1       
        if isWordTypeTable(child):
            tableObject[child.tag] = parseTable(child, table.cell(zeile,spalte))
        elif isWordTypeCheckboxlist(child):
            tableObject[child.tag] = parseCheckboxlist(child, table.cell(zeile,spalte))
        elif isWordTypeYesNolist(child):
            tableObject[child.tag] = parseYesNoList(child, table.cell(zeile,spalte))
        else:
            tableObject[child.tag] = table.cell(zeile,spalte).text      

    return tableObject


def parseChecklist(checklistFile):
    log = logging.getLogger("DSEGenerator.checklistParser")
    try:
        checklistDocument = Document(checklistFile)
    except (FileNotFoundError):
        log.error("File '" + checklistFile + "' not Found! " + FileNotFoundError.strerror)
        checklistDocument = None

    if checklistDocument is not None:
        version = readVersion(checklistDocument)
        checklistTemplate = Resources.getChecklisteTemplate(version)

        try:
            tree = ET.parse(checklistTemplate) 
        except (ET.ParseError):
            log.error("Checklist Template '" + checklistTemplate + "' could not be read!")
            tree = None
            checklistObject = None
        except (FileNotFoundError):
            wx.MessageBox("Error occured when opening Checklist Template! " + FileNotFoundError.strerror, caption="Error occured!")
            log.error("Checklist Template '" + checklistTemplate + "' could not be found! " + FileNotFoundError.strerror)

        if tree != None:
            root = tree.getroot()
            checklistObject = XMLObject()
            checklistObject.xmlVersion = root.attrib.get(const.DSEDOC_ATTRIB_VERSION)
            if Resources.validVersions(version, checklistObject.xmlVersion):
                for elem in root:
                    if isWordTypeTable(elem):
                        checklistObject.addElement(elem.tag, parseTable(elem, checklistDocument))

                if const.CHECKLIST_ATTRIB_TITLE in checklistObject.elementList:
                    checklistObject.wordVersion = checklistObject.elementList[const.CHECKLIST_ATTRIB_TITLE][const.CHECKLIST_ATTRIB_VERSION]
            else:
                log.warn("No valid versions! Processing skipped!")
    else:
        log.error("Processing of Template skipped! Please check Error log!")
        checklistObject = None
    log.debug("finished creation checklistObject: " + str(checklistObject))
    return checklistObject