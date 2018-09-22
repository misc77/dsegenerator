import const
import logger
import os
import xml.etree.ElementTree as ET
from xmlObject import XMLObject
from docx import Document
from resources import Resources

class DocGenerator:
   
    def __init__(self, checklist):
        self.checklistObject = checklist
        self.dseTemplate = {}
        self.dseDocument = None
        self.processed = False


    def evaluateFormular(self, text):
        log = logger.getLogger()
        startpos = text.index("{")
        if startpos > 0:
            endpos = text.index("}")+1
            if endpos > startpos:
                skript = text[startpos:endpos]
                cleanSkript = skript.replace("{","").replace("}","")
                try:
                    text = text.replace(skript, eval(cleanSkript)) 
                    log.info("Evaluated Text: " + text)
                except (IndexError, OverflowError, SyntaxError, TypeError, NameError):                    
                        log.warn("Error occured! Skipped evaluation! String to evaluate '" + text + "' will be replaced by ''.")
                        text = text.replace(skript, "")    

                if text.count("{") > 0:
                    text = self.evaluateFormular(text)
            else:
                log.warn("Skipped processing! End position of symbol of text to evaluate is before start position (start:" + startpos +", end:" + endpos+")")
        return text


    def processParagraph(self, paragraph):
        titel = paragraph.attrib.get(const.DSEDOC_ATTRIB_TITLE)
        if titel != "":
            self.dseDocument.add_heading(titel, level=3)
        text = paragraph.text
        if text != "":
            if text.count("{") > 0:    
                text = self.evaluateFormular(text)                
            self.dseDocument.add_paragraph(text)


    def processChapter(self, chapter):
        titel = chapter.attrib.get(const.DSEDOC_ATTRIB_TITLE)
        if titel != "":
            self.dseDocument.add_heading(titel, level=2)
        for elem in chapter:
            if elem.tag == const.DSEDOC_TAG_PARAGRAPH:
                self.processParagraph(elem)


    def parseTemplate(self, version = "1.0"):
        log = logger.getLoggerCtx("DSEGenerator.docGenerator.parseTemplate")
        filename = Resources.getDSETemplate(version)
        try:
            tree = ET.parse(filename)   
        except(ET.ParseError):
            tree = None
            log.error("Error occured when parsing XML document '" + filename + "'! " + ET.ParseError.text)
        if tree is not None:
            root = tree.getroot()
            if Resources.validVersions(self.checklistObject.xmlVersion, root.attrib.get(const.DSEDOC_ATTRIB_VERSION)):
                self.dseDocument = Document()
                core_properties = self.dseDocument.core_properties
                core_properties.comments = "Checklist Version:" + self.checklistObject.wordVersion + ", Checklist Template Version: " + self.checklistObject.xmlVersion + ", DSE Document Template Version: " + root.attrib.get(const.DSEDOC_ATTRIB_VERSION)
                for elem in root:
                    if elem.tag == const.DSEDOC_TAG_CHAPTER:
                        self.processChapter(elem)
                self.processed = True
            else:
                log.warn("Processing skipped because of invalid versions between Checklist template XML and DSE template XML!! ")


    def saveDocument(self, versionnumber=1, path=None):
        log = logger.getLoggerCtx("DSEGenerator.docGenerator.saveDocument")
        if path is None:
            filename = Resources.getOutputPath() + "/" + self.checklistObject.created.strftime("%Y%m%d%H%M%S") + "_dseDocument_"+str(versionnumber)+".docx"  
        else:
            filename = path      
        try:
            self.dseDocument.save(filename)
        except (PermissionError):
            log.warn("File '" + filename + "' could not be written! " + PermissionError.strerror)
            self.saveDocument(versionnumber+1)

        if os.path.isfile(filename):
            log.info("File '" + filename + "' has been written successfully!")
            return True
        else:
            log.warn("File '" + filename + "' has NOT been written! Please check error log!")
            return False


    